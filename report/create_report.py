from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont
import pandas as pd

BASE = Path('/mnt/data/sre_assignment6_package')
REPORT_DIR = BASE / 'report'
REPORT_DIR.mkdir(exist_ok=True)

def make_placeholder(path: Path, title: str, subtitle: str):
    img = Image.new('RGB', (1200, 560), (245, 245, 245))
    draw = ImageDraw.Draw(img)
    draw.rectangle((20, 20, 1180, 540), outline=(90, 90, 90), width=4)
    try:
        font_title = ImageFont.truetype('/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf', 44)
        font_sub = ImageFont.truetype('/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf', 30)
    except:
        font_title = font_sub = None
    draw.text((60, 190), title, fill=(40, 40, 40), font=font_title)
    draw.text((60, 270), subtitle, fill=(80, 80, 80), font=font_sub)
    img.save(path)

make_placeholder(REPORT_DIR/'placeholder_hpa.png', 'INSERT SCREENSHOT HERE', 'kubectl get hpa with 80% CPU target')
make_placeholder(REPORT_DIR/'placeholder_locust.png', 'INSERT SCREENSHOT HERE', 'Locust load test output / web UI')
make_placeholder(REPORT_DIR/'placeholder_pods.png', 'INSERT SCREENSHOT HERE', 'kubectl get pods showing scaled replicas')

csv = pd.read_csv(BASE/'outputs'/'monthly_traffic_forecast.csv')

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False):
    cell.text = ''
    p = cell.paragraphs[0]
    r = p.add_run(str(text))
    r.bold = bold
    r.font.size = Pt(9)

def add_code(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    for line in text.strip().split('\n'):
        run = p.add_run(line + '\n')
        run.font.name = 'Courier New'
        run.font.size = Pt(8)


doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.7)
sec.bottom_margin = Inches(0.7)
sec.left_margin = Inches(0.8)
sec.right_margin = Inches(0.8)

styles = doc.styles
styles['Normal'].font.name = 'Calibri'
styles['Normal'].font.size = Pt(10)
styles['Heading 1'].font.size = Pt(16)
styles['Heading 2'].font.size = Pt(13)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Assignment 6: Automation in SRE, Capacity Planning and Load Testing')
r.bold = True
r.font.size = Pt(18)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Individual Project Report')
r.font.size = Pt(13)

for label in ['Student: ____________________', 'Course: Site Reliability Engineering', 'Date: May 2026']:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(label)

doc.add_page_break()

doc.add_heading('1. Introduction', level=1)
doc.add_paragraph(
    'This report demonstrates automation in Site Reliability Engineering for a growing e-commerce API. '
    'The work includes automated log extraction, monthly traffic growth calculation, six-month traffic forecasting, '
    'Kubernetes Horizontal Pod Autoscaler configuration, traffic spike simulation using Locust, and ROI/incident analysis.'
)

doc.add_heading('2. Objective and Scenario', level=1)
doc.add_paragraph(
    'The objective is to understand how automation helps SRE teams predict traffic growth, prepare infrastructure capacity, '
    'and keep the service stable during unexpected load spikes. The scenario is a rapidly growing e-commerce API service.'
)

doc.add_heading('3. Predictive Analysis: Python Log Extraction and Forecasting', level=1)
doc.add_paragraph(
    'The Python script analysis/traffic_forecast.py automatically extracts access log records, groups requests by month, '
    'calculates month-to-month growth, and forecasts the next six months using average monthly growth.'
)

doc.add_heading('3.1 Forecasting Logic', level=2)
doc.add_paragraph(
    'The formula used for each future month is: forecast = last_actual_traffic * (1 + average_growth)^n, '
    'where n is the number of months after the latest actual month. In this dataset, the calculated average monthly growth is 17.94%.'
)

doc.add_heading('3.2 Forecast Output', level=2)
table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
for i, title in enumerate(['Month', 'Actual Requests', 'Forecast Requests']):
    set_cell_text(hdr[i], title, True)
    set_cell_shading(hdr[i], 'D9EAF7')
for _, row in csv.iterrows():
    cells = table.add_row().cells
    set_cell_text(cells[0], row['month'])
    set_cell_text(cells[1], '' if pd.isna(row['actual_requests']) else int(row['actual_requests']))
    set_cell_text(cells[2], '' if pd.isna(row['forecast_requests']) else int(row['forecast_requests']))
for row in table.rows:
    for cell in row.cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

doc.add_paragraph('Figure 1. Forecasted traffic output generated by Python.')
doc.add_picture(str(BASE/'outputs'/'traffic_forecast.png'), width=Inches(6.3))

doc.add_heading('4. Infrastructure Scaling with Kubernetes HPA', level=1)
doc.add_paragraph(
    'The e-commerce API is deployed as a Kubernetes Deployment. CPU requests and limits are configured because HPA CPU utilization is calculated relative to requested CPU. '
    'The Horizontal Pod Autoscaler targets the Deployment and keeps average CPU utilization around 80% by increasing or decreasing the number of pods.'
)

doc.add_heading('4.1 HPA Configuration', level=2)
add_code(doc, (BASE/'k8s'/'hpa.yaml').read_text())

doc.add_paragraph('Figure 2. HPA status after applying hpa.yaml.')
doc.add_picture(str(REPORT_DIR/'placeholder_hpa.png'), width=Inches(6.3))

doc.add_heading('5. Load Testing Simulation', level=1)
doc.add_paragraph(
    'Locust is used to simulate traffic spikes. The test sends frequent requests to /cpu?ms=250, an intentionally CPU-heavy endpoint. '
    'This increases pod CPU usage and allows HPA to demonstrate automatic scaling.'
)

doc.add_heading('5.1 Load Test Command', level=2)
add_code(doc, 'locust -f load_test/locustfile.py --host http://localhost:8080 --users 150 --spawn-rate 20 --run-time 3m --headless')

doc.add_paragraph('Figure 3. Locust load testing output.')
doc.add_picture(str(REPORT_DIR/'placeholder_locust.png'), width=Inches(6.3))

doc.add_paragraph('Figure 4. Kubernetes pods after HPA scaling.')
doc.add_picture(str(REPORT_DIR/'placeholder_pods.png'), width=Inches(6.3))

doc.add_heading('6. Real-World Incidents Preventable by Automation', level=1)
items = [
    ('Traffic spike overload', 'Without autoscaling, sudden campaigns or seasonal sales can overload a small number of pods. HPA prevents this by increasing replicas automatically when CPU usage rises.'),
    ('Slow manual capacity response', 'If engineers manually add capacity only after complaints start, users experience slow responses or failures. Automated monitoring and scaling reduce response time.'),
    ('Late detection of growth trends', 'If teams do not analyze logs, monthly traffic growth can be missed until the system is already unstable. Forecasting gives time to plan resources before incidents happen.'),
]
for title, body in items:
    p = doc.add_paragraph(style=None)
    p.add_run(title + ': ').bold = True
    p.add_run(body)

doc.add_heading('7. ROI Analysis of SRE Automation', level=1)
doc.add_paragraph(
    'SRE automation creates ROI by reducing downtime, manual engineering work, and over-provisioning costs. Forecasting helps predict future traffic and avoid last-minute infrastructure changes. '
    'HPA reduces wasted resources because the application can run with fewer pods during normal load and scale up only during traffic spikes. Load testing gives evidence that the system can survive higher demand before real users are affected.'
)

doc.add_paragraph(
    'Example ROI logic: if automation prevents even one one-hour outage during a high-traffic sale, the saved revenue and reputation can be higher than the cost of implementing scripts, monitoring, and autoscaling. '
    'Automation also improves engineer productivity because repeated tasks such as log extraction, forecast calculation, and scaling checks are handled by scripts and Kubernetes.'
)

doc.add_heading('8. Conclusion', level=1)
doc.add_paragraph(
    'This project shows how SRE automation connects prediction, scaling, and validation. Python forecasting identifies expected future load, Kubernetes HPA provides automatic scaling at an 80% CPU threshold, '
    'and Locust proves scaling behavior under simulated traffic. Together, these practices make the service more reliable and reduce operational risk.'
)

doc.add_heading('9. References', level=1)
for ref in [
    'Kubernetes Documentation: Horizontal Pod Autoscaling.',
    'Kubernetes Documentation: Horizontal Pod Autoscaler Walkthrough.',
    'Locust Documentation: Quickstart and headless load testing.',
    'Kubernetes Metrics Server Documentation.'
]:
    doc.add_paragraph(ref, style=None)

out = REPORT_DIR/'Assignment_6_Report.docx'
doc.save(out)
print(out)
